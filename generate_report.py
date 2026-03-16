import docx
from docx.shared import Pt, RGBColor
import os

def create_portfolio_report():
    doc = docx.Document()

    # Title
    title = doc.add_heading('포트폴리오: 율시스템 I3D 플랜트 유지보수 AI Copilot', 0)
    title.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

    # 1. Project Overview
    doc.add_heading('1. 프로젝트 개요', level=1)
    doc.add_paragraph(
        '본 프로젝트는 복잡한 산업용 플랜트 현장에서 발생하는 수많은 설비 경고와 유지보수 매뉴얼을, '
        'AI(LLM)가 스스로 확인하고 진단하여 조치까지 취해주는 «설비 유지보수 자율 Agent (Copilot)» 데모입니다.\n'
        '사내망(온프라미스/로컬) ERP 데이터베이스와 3D 디지털 트윈 뷰어 장비를 통합 제어하는 '
        'MCP(Model Context Protocol) 철학을 흉내 내어, 언어 모델이 로컬 시스템의 데이터 상태를 '
        '직접 조회하고 조작할 수 있도록 구현되었습니다.'
    )

    # 2. Key Features and Tech Stack
    doc.add_heading('2. 핵심 기능 및 기술 스택', level=1)
    
    doc.add_heading('주요 기능', level=2)
    features = [
        '대화형 LLM 에이전트 (ReAct 기반): 작업자의 자연어 명령("온도 높은 장비 찾아봐")을 분석하여 생각(Thought) -> 도구 선택(Action) -> 결과 관찰(Observation) 사이클을 자율적으로 수행합니다.',
        '사내 데이터 자율 연동: API 및 JSON 제어를 통해 가상 ERP와 3D 디지털 트윈 뷰어에 연동됩니다.',
        '지식 검색 엔진 (Hybrid RAG): 정규표현식을 결합한 하이브리드 재정렬(Re-ranking)을 통해 장비 번호 오인식과 같은 고유명사 매칭의 한계를 극복하고 100%의 정확성을 목표로 합니다.',
        '프리미엄 대시보드 UI (Streamlit): LLM이 데이터를 조작하면 실시간 대시보드가 리로딩되어 실제 관제 화면과 같은 UI를 제공합니다.'
    ]
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_heading('기술 스택', level=2)
    techs = [
        'Core AI Engine: Google Gemini 2.5 Pro, LangGraph',
        'RAG Pipeline: ChromaDB, HuggingFace (jhgan/ko-sroberta-multitask)',
        'Frontend & Data: Streamlit, Json Mock DB'
    ]
    for tech in techs:
        doc.add_paragraph(tech, style='List Bullet')

    # 3. Core Scenarios (Placeholders for screenshots)
    doc.add_heading('3. 핵심 유스케이스 및 시연 (시나리오 테스트)', level=1)

    # Scenario 1
    doc.add_heading('상황 1: 설비 매뉴얼 검색 및 가이드 제공', level=2)
    doc.add_paragraph('Prompt: "v-103 메뉴얼 알려줘"')
    doc.add_paragraph(
        '설명: 수십 페이지 분량의 방대한 매뉴얼 텍스트 파일 속에서 V-103 장비에 대한 '
        '유지보수 대응법을 RAG 검색 엔진으로 정확하게 추출하여 답변을 제공합니다.'
    )
    p1 = doc.add_paragraph()
    p1.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run('\n[여기에 V-103 메뉴얼 검색 캡쳐화면 삽입]\n')
    r1.font.color.rgb = RGBColor(255, 0, 0)
    r1.font.bold = True

    # Scenario 2
    doc.add_heading('상황 2: 비정상 상태 감지 및 자율 조치 수행', level=2)
    doc.add_paragraph('Prompt: "정상상태가 아니니 조치취해줘"')
    doc.add_paragraph(
        '설명: 에이전트가 직접 시스템 상태(ERP Data)를 조회해 비정상 설비를 감지하고, '
        '제공된 Tool(Tools Usage)을 활용해 장비 파라미터를 정상 상태로 업데이트하며 '
        '3D 뷰어 카메라를 해당 장비로 자동 이동(포커싱)시킵니다. 데이터가 변경되면 '
        'UI 대시보드 역시 실시간으로 연동되어 리런(rerun)됩니다.'
    )
    p2 = doc.add_paragraph()
    p2.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run('\n[여기에 조치 수행 및 UI 상태 변화 캡쳐화면 삽입]\n')
    r2.font.color.rgb = RGBColor(255, 0, 0)
    r2.font.bold = True

    # Scenario 3
    doc.add_heading('상황 3: 환각(Hallucination) 방지 및 안전한 대응', level=2)
    doc.add_paragraph('Prompt: "(예시) 잘못된 정보나 매뉴얼에 없는 내용 질문"')
    doc.add_paragraph(
        '설명: 잘못된 정보나 DB에 없는 내용에 대해서는 억측하지 않고, '
        '"모른다"고 명확히 답변하도록 설계되어 산업 현장의 시스템 안정성 및 무결성을 보장합니다.'
    )
    p3 = doc.add_paragraph()
    p3.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run('\n[여기에 모른다고 명확히 답변하는 캡쳐화면 삽입]\n')
    r3.font.color.rgb = RGBColor(255, 0, 0)
    r3.font.bold = True

    # 4. Troubleshooting
    doc.add_heading('4. 트러블슈팅 및 고도화 아키텍처', level=1)
    doc.add_paragraph(
        '단순 튜토리얼 수준을 넘어 LLM의 산업 도입 시 발생하는 여러 한계점을 '
        '극복하기 위해 아래와 같은 트러블슈팅 및 고도화를 진행했습니다.'
    )
    troubleshoots = [
        '인공지능 컨텍스트 기억 상실 에러 복구 및 Memory Management 구축',
        '명시적 가드레일 부재로 인한 환각 현상 분석 및 안전장치 도입',
        'Dense Retrieval(Vector DB) 망각 및 중첩 문제 해결',
        '정규식(Regex)을 이용한 2차 하이브리드 리랭킹 도입을 통해 고유 식별 번호 기반 검색 정확도 확보'
    ]
    for ts in troubleshoots:
        doc.add_paragraph(ts, style='List Bullet')

    # Save the document
    output_path = r'C:\YoulSystem\Portfolio_AI_Copilot.docx'
    doc.save(output_path)
    print(f"Report successfully saved to: {output_path}")

if __name__ == '__main__':
    create_portfolio_report()
